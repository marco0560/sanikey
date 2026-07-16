"""Document inventory and text extraction."""

from __future__ import annotations

import fnmatch
import hashlib
import re
import shutil
import subprocess
import tarfile
import tempfile
import warnings
import zipfile
from dataclasses import dataclass
from pathlib import Path, PurePath
from typing import TYPE_CHECKING

from .models import DocumentRecord

if TYPE_CHECKING:
    from .config import IngestionConfig, PersonConfig
    from .progress import ProgressReporter

DATE_PREFIX_RE = re.compile(r"^(?P<date>\d{8})[\s_-]+(?P<title>.+)$")
DICOM_EXTENSIONS = {".dcm": "dicom_file", ".img": "dicom_img", ".iso": "dicom_iso"}
TEXT_EXTENSIONS = {".txt", ".md"}
IMAGE_EXTENSIONS = {".jpeg", ".jpg", ".png"}
ARCHIVE_EXTENSIONS = {".7z", ".rar", ".tar.xz", ".zip"}
DOCX_EXTENSIONS = {".docx"}
LEGACY_OFFICE_EXTENSIONS = {".doc"}
ODF_EXTENSIONS = {".odt"}
OFFICE_EXTENSIONS = DOCX_EXTENSIONS | LEGACY_OFFICE_EXTENSIONS | ODF_EXTENSIONS
SPREADSHEET_EXTENSIONS = {".ods", ".xls", ".xlsb", ".xlsm", ".xlsx"}
MIN_PDF_TEXT_CHARACTERS = 40


@dataclass(frozen=True)
class ExtractedText:
    """Text extracted from one document.

    Parameters
    ----------
    document_id : str
        Source document id.
    text : str
        Extracted text content.
    warnings : tuple[str, ...]
        Non-fatal extraction warnings.
    """

    document_id: str
    text: str
    warnings: tuple[str, ...] = ()


@dataclass(frozen=True)
class DocumentRecordOrigin:
    """Provenance fields for a document record.

    Parameters
    ----------
    origin : str
        Document origin, either ``source`` or ``container``.
    container_id : str | None, optional
        Parent container document id.
    internal_path : str | None, optional
        Member path inside the parent container.
    """

    origin: str = "source"
    container_id: str | None = None
    internal_path: str | None = None


def scan_documents(person: PersonConfig) -> tuple[DocumentRecord, ...]:
    """Scan a patient's source document directory.

    Parameters
    ----------
    person : PersonConfig
        Patient configuration.

    Returns
    -------
    tuple[DocumentRecord, ...]
        Deterministically sorted document records.
    """

    return _deduplicate_documents(scan_document_inventory(person))


def scan_document_inventory(
    person: PersonConfig,
    *,
    progress: ProgressReporter | None = None,
    progress_label: str | None = None,
) -> tuple[DocumentRecord, ...]:
    """Scan every file in a patient's source document directory.

    Parameters
    ----------
    person : PersonConfig
        Patient configuration.
    progress : ProgressReporter | None, optional
        Progress reporter for scanned files.
    progress_label : str | None, optional
        Label for the progress line.

    Returns
    -------
    tuple[DocumentRecord, ...]
        Deterministically sorted document records, including duplicate content.
    """

    root = person.source_documents
    if not root.exists():
        return ()
    files = sorted(
        path
        for path in root.rglob("*")
        if path.is_file() and not is_excluded_source_path(person, path)
    )
    if progress is not None and progress_label is not None:
        progress.begin(progress_label, total=len(files), interval=20)
    records = []
    for index, path in enumerate(files, start=1):
        records.append(document_record_for_path(path, root=root, patient_id=person.id))
        if progress is not None:
            progress.advance(index, total=len(files))
    if progress is not None and progress_label is not None:
        progress.done(f"done files={len(files)}")
    return tuple(records)


def excluded_source_files(person: PersonConfig) -> tuple[Path, ...]:
    """Return source files excluded by ingestion configuration.

    Parameters
    ----------
    person : PersonConfig
        Patient configuration.

    Returns
    -------
    tuple[pathlib.Path, ...]
        Deterministically sorted excluded source files.
    """

    root = person.source_documents
    if not root.exists():
        return ()
    return tuple(
        sorted(
            path
            for path in root.rglob("*")
            if path.is_file() and is_excluded_source_path(person, path)
        )
    )


def is_excluded_source_path(person: PersonConfig, path: Path) -> bool:
    """Return whether a source path matches ingestion exclusions.

    Parameters
    ----------
    person : PersonConfig
        Patient configuration.
    path : pathlib.Path
        Candidate source file.

    Returns
    -------
    bool
        ``True`` when the path should be skipped by ingestion.
    """

    try:
        relative = path.relative_to(person.source_documents).as_posix()
    except ValueError:
        relative = path.name
    return _is_excluded_by_ingestion(relative, person.ingestion)


def _is_excluded_by_ingestion(relative_path: str, ingestion: IngestionConfig) -> bool:
    """Return whether ingestion rules exclude a relative path.

    Parameters
    ----------
    relative_path : str
        POSIX-style path relative to a source or container root.
    ingestion : IngestionConfig
        Include and exclude glob patterns.

    Returns
    -------
    bool
        ``True`` when an exclusion matches and no inclusion recovers the path.
    """

    if _matches_patterns(relative_path, ingestion.include_patterns):
        return False
    return _matches_patterns(relative_path, ingestion.exclude_patterns)


def _matches_patterns(relative_path: str, patterns: tuple[str, ...]) -> bool:
    """Return whether a relative path matches any configured pattern.

    Parameters
    ----------
    relative_path : str
        POSIX-style path relative to a source or container root.
    patterns : tuple[str, ...]
        Glob patterns to evaluate.

    Returns
    -------
    bool
        ``True`` when any pattern matches the path or filename. Matching is
        case-insensitive to tolerate inconsistent directory names from clinical
        media.
    """

    normalized_path = relative_path.casefold()
    name = PurePath(normalized_path).name
    return any(
        _matches_pattern(normalized_path, name, pattern.casefold())
        for pattern in patterns
    )


def _matches_pattern(normalized_path: str, name: str, pattern: str) -> bool:
    """Return whether one normalized path matches one normalized glob.

    Parameters
    ----------
    normalized_path : str
        POSIX-style case-folded path.
    name : str
        Case-folded basename.
    pattern : str
        Case-folded glob pattern.

    Returns
    -------
    bool
        ``True`` when the pattern matches the full path, basename, or a root
        path covered by a leading ``**/`` glob.
    """

    candidates = [pattern]
    if pattern.startswith("**/"):
        candidates.append(pattern[3:])
    return any(
        fnmatch.fnmatchcase(normalized_path, candidate)
        or fnmatch.fnmatchcase(name, candidate)
        for candidate in candidates
    )


def find_duplicate_documents(
    documents: tuple[DocumentRecord, ...],
) -> dict[str, tuple[DocumentRecord, ...]]:
    """Find technical duplicate documents by digest.

    Parameters
    ----------
    documents : tuple[DocumentRecord, ...]
        Document records to inspect.

    Returns
    -------
    dict[str, tuple[DocumentRecord, ...]]
        SHA256 digests with more than one document.
    """

    grouped: dict[str, list[DocumentRecord]] = {}
    for document in documents:
        grouped.setdefault(document.sha256, []).append(document)
    return {digest: tuple(items) for digest, items in grouped.items() if len(items) > 1}


def duplicate_document_warnings(
    duplicates: dict[str, tuple[DocumentRecord, ...]],
) -> tuple[str, ...]:
    """Build warnings for duplicate-content documents.

    Parameters
    ----------
    duplicates : dict[str, tuple[DocumentRecord, ...]]
        Duplicate digest groups.

    Returns
    -------
    tuple[str, ...]
        Human-readable warnings naming the identical files.
    """

    warnings: list[str] = []
    for digest, documents in sorted(duplicates.items()):
        retained = documents[0]
        for skipped in documents[1:]:
            warnings.append(
                "contenuto documento duplicato saltato. "
                "I file seguenti sono identici "
                f"(sha256={digest}): \n{retained.path}\n{skipped.path}"
            )
    return tuple(warnings)


def extract_text(document: DocumentRecord) -> ExtractedText:
    """Extract text from a document when supported.

    Parameters
    ----------
    document : DocumentRecord
        Source document.

    Returns
    -------
    ExtractedText
        Extracted text and non-fatal warnings.
    """

    suffix = _known_suffix(document.path)
    if suffix in TEXT_EXTENSIONS:
        return ExtractedText(
            document_id=document.document_id,
            text=document.path.read_text(encoding="utf-8", errors="replace"),
        )
    if suffix == ".pdf":
        return _extract_pdf_text(document)
    if suffix in ARCHIVE_EXTENSIONS:
        return _extract_archive_text(document)
    if suffix in DOCX_EXTENSIONS:
        return _extract_docx_text(document)
    if suffix in SPREADSHEET_EXTENSIONS:
        return _extract_spreadsheet_text(document)
    if suffix in ODF_EXTENSIONS:
        return _extract_odf_text(document)
    if suffix in LEGACY_OFFICE_EXTENSIONS:
        return _extract_legacy_office_text(document)
    if suffix in IMAGE_EXTENSIONS:
        return ExtractedText(document_id=document.document_id, text="")
    if document.kind.startswith("dicom_"):
        return ExtractedText(
            document_id=document.document_id,
            text="",
        )
    return ExtractedText(
        document_id=document.document_id,
        text="",
        warnings=(
            f"estrazione testo non supportata per {suffix or 'file senza estensione'}",
        ),
    )


def document_page_count(document: DocumentRecord) -> int | None:
    """Return the document page count when it can be determined reliably.

    Parameters
    ----------
    document : DocumentRecord
        Source document.

    Returns
    -------
    int | None
        Page count, or ``None`` when the format does not expose a reliable page
        count in the local extractor.
    """

    suffix = _known_suffix(document.path)
    if suffix in TEXT_EXTENSIONS:
        return 1
    if suffix == ".pdf":
        return _pdf_page_count(document.path)
    return None


def document_record_for_path(
    path: Path,
    *,
    root: Path,
    patient_id: str,
    provenance: DocumentRecordOrigin | None = None,
) -> DocumentRecord:
    """Build a document record for one file.

    Parameters
    ----------
    path : pathlib.Path
        Source file.
    root : pathlib.Path
        Source document root.
    patient_id : str
        Owning patient id.
    provenance : DocumentRecordOrigin | None, optional
        Optional provenance fields.

    Returns
    -------
    DocumentRecord
        Document record.
    """

    digest = _sha256(path)
    record_origin = provenance or DocumentRecordOrigin()
    relative = path.relative_to(root)
    category = _document_category(relative)
    date, title = _title_from_name(_document_title_stem(path))
    kind = _document_kind(path)
    document_id = digest
    if record_origin.origin != "source":
        identity = "\0".join(
            (
                record_origin.origin,
                record_origin.container_id or "",
                record_origin.internal_path or "",
                digest,
            )
        )
        document_id = hashlib.sha256(identity.encode("utf-8")).hexdigest()
    return DocumentRecord(
        document_id=document_id,
        patient_id=patient_id,
        path=path,
        title=title,
        category=category,
        kind=kind,
        sha256=digest,
        date=date,
        origin=record_origin.origin,
        container_id=record_origin.container_id,
        internal_path=record_origin.internal_path,
    )


def _document_category(relative: Path) -> str:
    """Return the display category for a source-relative document path.

    Parameters
    ----------
    relative : pathlib.Path
        Source-relative document path.

    Returns
    -------
    str
        Top-level category name. A leading underscore marks service
        directories for sorting and is not part of the display category.
    """

    if len(relative.parts) <= 1:
        return "uncategorized"
    return relative.parts[0].lstrip("_") or "uncategorized"


def _deduplicate_documents(
    documents: tuple[DocumentRecord, ...],
) -> tuple[DocumentRecord, ...]:
    """Keep only the first document for each content digest.

    Parameters
    ----------
    documents : tuple[DocumentRecord, ...]
        Scanned document records in deterministic order.
    Returns
    -------
    tuple[DocumentRecord, ...]
        First document for each SHA256 digest.
    """

    seen: set[str] = set()
    retained: list[DocumentRecord] = []
    for document in documents:
        if document.sha256 in seen:
            continue
        seen.add(document.sha256)
        retained.append(document)
    return tuple(retained)


def _sha256(path: Path) -> str:
    """Compute a file SHA256 digest.

    Parameters
    ----------
    path : pathlib.Path
        File to hash.

    Returns
    -------
    str
        Hex digest.
    """

    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _title_from_name(stem: str) -> tuple[str | None, str]:
    """Extract optional date and title from a filename stem.

    Parameters
    ----------
    stem : str
        Filename stem.

    Returns
    -------
    tuple[str | None, str]
        ISO date when present and a readable title.
    """

    match = DATE_PREFIX_RE.match(stem)
    if match is None:
        return None, stem.replace("_", " ").strip()
    raw_date = match.group("date")
    date = f"{raw_date[:4]}-{raw_date[4:6]}-{raw_date[6:]}"
    return date, match.group("title").replace("_", " ").strip()


def _document_kind(path: Path) -> str:
    """Return document kind from extension.

    Parameters
    ----------
    path : pathlib.Path
        Source file.

    Returns
    -------
    str
        Document kind.
    """

    suffix = _known_suffix(path)
    if _has_dicom_magic(path):
        return "dicom_file"
    if suffix in DICOM_EXTENSIONS:
        return DICOM_EXTENSIONS[suffix]
    if suffix == ".pdf":
        return "pdf"
    if suffix in TEXT_EXTENSIONS:
        return "text"
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    if suffix in ARCHIVE_EXTENSIONS:
        return "archive"
    if suffix in OFFICE_EXTENSIONS | SPREADSHEET_EXTENSIONS:
        return "office"
    return "binary"


def _known_suffix(path: Path) -> str:
    """Return the known logical suffix for a path.

    Parameters
    ----------
    path : pathlib.Path
        Path to inspect.

    Returns
    -------
    str
        Multi-part suffix such as ``.tar.xz`` when recognized, otherwise the
        final path suffix.
    """

    suffixes = tuple(item.lower() for item in path.suffixes)
    if suffixes[-2:] == (".tar", ".xz"):
        return ".tar.xz"
    return path.suffix.lower()


def _document_title_stem(path: Path) -> str:
    """Return the filename stem used for document titles.

    Parameters
    ----------
    path : pathlib.Path
        Path to inspect.

    Returns
    -------
    str
        Filename without a recognized suffix.
    """

    suffix = _known_suffix(path)
    if suffix and path.name.lower().endswith(suffix):
        return path.name[: -len(suffix)]
    return path.stem


def _has_dicom_magic(path: Path) -> bool:
    """Return whether a file contains the standard DICOM magic marker.

    Parameters
    ----------
    path : pathlib.Path
        File to inspect.

    Returns
    -------
    bool
        ``True`` when bytes 128-131 are ``DICM``.
    """

    try:
        with path.open("rb") as handle:
            handle.seek(128)
            return handle.read(4) == b"DICM"
    except OSError:
        return False


def _extract_pdf_text(document: DocumentRecord) -> ExtractedText:
    """Extract PDF text through the first available provider.

    Parameters
    ----------
    document : DocumentRecord
        PDF document.

    Returns
    -------
    ExtractedText
        Extracted text or warning when no provider is available.
    """

    pymupdf_result = _extract_pdf_text_with_pymupdf(document)
    if pymupdf_result is not None and _has_sufficient_pdf_text(pymupdf_result.text):
        return pymupdf_result
    ocrmypdf_result = _extract_pdf_text_with_ocrmypdf(document)
    if ocrmypdf_result is not None:
        if pymupdf_result is not None and pymupdf_result.warnings:
            return ExtractedText(
                document_id=document.document_id,
                text=ocrmypdf_result.text,
                warnings=(*pymupdf_result.warnings, *ocrmypdf_result.warnings),
            )
        return ocrmypdf_result
    if pymupdf_result is not None:
        if pymupdf_result.warnings:
            return pymupdf_result
        return ExtractedText(
            document_id=document.document_id,
            text=pymupdf_result.text,
            warnings=(
                "PyMuPDF ha estratto testo insufficiente e non e' disponibile "
                "un provider OCR",
            ),
        )
    return ExtractedText(
        document_id=document.document_id,
        text="",
        warnings=(
            "Nessun provider di estrazione testo PDF disponibile; installare PyMuPDF "
            "o configurare OCRmyPDF",
        ),
    )


def _extract_pdf_text_with_pymupdf(document: DocumentRecord) -> ExtractedText | None:
    """Extract PDF text with PyMuPDF when installed.

    Parameters
    ----------
    document : DocumentRecord
        PDF document.

    Returns
    -------
    ExtractedText | None
        Extracted text, or ``None`` when PyMuPDF is unavailable.
    """

    try:
        import fitz
    except ImportError:
        return None
    display_errors = bool(fitz.TOOLS.mupdf_display_errors())
    display_warnings = bool(fitz.TOOLS.mupdf_display_warnings())
    try:
        fitz.TOOLS.mupdf_display_errors(False)
        fitz.TOOLS.mupdf_display_warnings(False)
        with fitz.open(document.path) as pdf:
            text = "\n".join(page.get_text() for page in pdf)
    except (fitz.FileDataError, RuntimeError, ValueError) as exc:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=(
                "PyMuPDF non ha potuto estrarre il testo PDF; uso OCRmyPDF "
                f"se disponibile: {exc}",
            ),
        )
    finally:
        fitz.TOOLS.mupdf_display_errors(display_errors)
        fitz.TOOLS.mupdf_display_warnings(display_warnings)
    return ExtractedText(document_id=document.document_id, text=text)


def _has_sufficient_pdf_text(text: str) -> bool:
    """Return whether extracted PDF text is sufficient to skip OCR.

    Parameters
    ----------
    text : str
        Extracted text.

    Returns
    -------
    bool
        ``True`` when text has enough non-whitespace characters.
    """

    return sum(1 for character in text if not character.isspace()) >= (
        MIN_PDF_TEXT_CHARACTERS
    )


def _extract_archive_text(document: DocumentRecord) -> ExtractedText:
    """Extract an archive member inventory.

    Parameters
    ----------
    document : DocumentRecord
        Archive document.

    Returns
    -------
    ExtractedText
        Textual archive inventory or warning.
    """

    suffix = _known_suffix(document.path)
    if suffix == ".zip":
        return _extract_zip_inventory(document)
    if suffix == ".7z":
        return _extract_7z_inventory(document)
    if suffix == ".rar":
        return _extract_rar_inventory(document)
    if suffix == ".tar.xz":
        return _extract_tar_xz_inventory(document)
    return ExtractedText(
        document_id=document.document_id,
        text="",
        warnings=(f"formato archivio non supportato {suffix}",),
    )


def _extract_zip_inventory(document: DocumentRecord) -> ExtractedText:
    """Extract ZIP member names.

    Parameters
    ----------
    document : DocumentRecord
        ZIP document.

    Returns
    -------
    ExtractedText
        Textual inventory or warning.
    """

    try:
        with zipfile.ZipFile(document.path) as archive:
            names = archive.namelist()
    except (OSError, zipfile.BadZipFile) as exc:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=(f"estrazione inventario ZIP non riuscita: {exc}",),
        )
    return ExtractedText(
        document_id=document.document_id,
        text=_archive_inventory_text("zip", names),
    )


def _extract_7z_inventory(document: DocumentRecord) -> ExtractedText:
    """Extract 7z member names.

    Parameters
    ----------
    document : DocumentRecord
        7z document.

    Returns
    -------
    ExtractedText
        Textual inventory or warning.
    """

    try:
        import py7zr

        with py7zr.SevenZipFile(document.path) as archive:
            names = archive.getnames()
    except (OSError, py7zr.Bad7zFile, py7zr.PasswordRequired) as exc:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=(f"estrazione inventario 7z non riuscita: {exc}",),
        )
    return ExtractedText(
        document_id=document.document_id,
        text=_archive_inventory_text("7z", names),
    )


def _extract_rar_inventory(document: DocumentRecord) -> ExtractedText:
    """Extract RAR member names.

    Parameters
    ----------
    document : DocumentRecord
        RAR document.

    Returns
    -------
    ExtractedText
        Textual inventory or warning.
    """

    try:
        import rarfile

        with rarfile.RarFile(document.path) as archive:
            names = archive.namelist()
    except (OSError, rarfile.Error) as exc:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=(f"estrazione inventario RAR non riuscita: {exc}",),
        )
    return ExtractedText(
        document_id=document.document_id,
        text=_archive_inventory_text("rar", names),
    )


def _extract_tar_xz_inventory(document: DocumentRecord) -> ExtractedText:
    """Extract TAR.XZ member names.

    Parameters
    ----------
    document : DocumentRecord
        TAR.XZ document.

    Returns
    -------
    ExtractedText
        Textual inventory or warning.
    """

    try:
        with tarfile.open(document.path, mode="r:xz") as archive:
            names = [member.name for member in archive.getmembers() if member.isfile()]
    except (OSError, tarfile.TarError) as exc:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=(f"estrazione inventario TAR.XZ non riuscita: {exc}",),
        )
    return ExtractedText(
        document_id=document.document_id,
        text=_archive_inventory_text("tar.xz", names),
    )


def _archive_inventory_text(kind: str, names: list[str]) -> str:
    """Render archive member names as text.

    Parameters
    ----------
    kind : str
        Archive kind.
    names : list[str]
        Member names.

    Returns
    -------
    str
        Textual inventory.
    """

    if not names:
        return f"archivio {kind} vuoto"
    return "\n".join((f"contenuto archivio {kind}:", *sorted(names)))


def _extract_docx_text(document: DocumentRecord) -> ExtractedText:
    """Extract text from a DOCX document.

    Parameters
    ----------
    document : DocumentRecord
        DOCX document.

    Returns
    -------
    ExtractedText
        Extracted text or warning.
    """

    try:
        import docx
        from docx.opc.exceptions import PackageNotFoundError

        parsed = docx.Document(str(document.path))
        parts = [paragraph.text for paragraph in parsed.paragraphs if paragraph.text]
        for table in parsed.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text)
    except (OSError, PackageNotFoundError, ValueError, zipfile.BadZipFile) as exc:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=(f"estrazione testo DOCX non riuscita: {exc}",),
        )
    return ExtractedText(document_id=document.document_id, text="\n".join(parts))


def _extract_spreadsheet_text(document: DocumentRecord) -> ExtractedText:
    """Extract text from a spreadsheet workbook.

    Parameters
    ----------
    document : DocumentRecord
        Spreadsheet document.

    Returns
    -------
    ExtractedText
        Extracted cell text or warning.
    """

    try:
        from python_calamine import CalamineError, load_workbook

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            workbook = load_workbook(document.path)
            try:
                parts = []
                for sheet_name in workbook.sheet_names:
                    worksheet = workbook.get_sheet_by_name(sheet_name)
                    parts.append(f"[{sheet_name}]")
                    for row in worksheet.to_python(skip_empty_area=False):
                        values = [
                            str(value) for value in row if value not in (None, "")
                        ]
                        if values:
                            parts.append("\t".join(values))
            finally:
                workbook.close()
    except (CalamineError, OSError, ValueError, zipfile.BadZipFile) as exc:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=(f"estrazione testo foglio di calcolo non riuscita: {exc}",),
        )
    warning_messages = tuple(
        dict.fromkeys(
            "testo foglio di calcolo estratto; funzionalita' di compatibilita' "
            "cartella non preservata: "
            f"{item.message}"
            for item in caught
        )
    )
    return ExtractedText(
        document_id=document.document_id,
        text="\n".join(parts),
        warnings=warning_messages,
    )


def _extract_odf_text(document: DocumentRecord) -> ExtractedText:
    """Extract text from ODF documents.

    Parameters
    ----------
    document : DocumentRecord
        ODT or ODS document.

    Returns
    -------
    ExtractedText
        Extracted text or warning.
    """

    try:
        from odf import table as odf_table
        from odf.opendocument import load
        from odf.teletype import extractText

        parsed = load(str(document.path))
        if document.path.suffix.lower() == ".ods":
            parts = []
            for sheet in parsed.spreadsheet.getElementsByType(odf_table.Table):
                sheet_name = sheet.getAttribute("name")
                if sheet_name:
                    parts.append(f"[{sheet_name}]")
                for row in sheet.getElementsByType(odf_table.TableRow):
                    values = [
                        extractText(cell)
                        for cell in row.getElementsByType(odf_table.TableCell)
                    ]
                    values = [value for value in values if value]
                    if values:
                        parts.append("\t".join(values))
            text = "\n".join(parts)
        else:
            text = extractText(parsed.text)
    except (AttributeError, OSError, ValueError, zipfile.BadZipFile) as exc:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=(f"estrazione testo ODF non riuscita: {exc}",),
        )
    return ExtractedText(document_id=document.document_id, text=text)


def _extract_legacy_office_text(document: DocumentRecord) -> ExtractedText:
    """Extract text from legacy Word documents through LibreOffice.

    Parameters
    ----------
    document : DocumentRecord
        Legacy ``.doc`` document.

    Returns
    -------
    ExtractedText
        Extracted text or warning.
    """

    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if executable is None:
        return ExtractedText(
            document_id=document.document_id,
            text="",
            warnings=("LibreOffice non installato; estrazione Office legacy saltata",),
        )
    with tempfile.TemporaryDirectory(prefix="sanikey-office-") as directory:
        output_dir = Path(directory)
        try:
            completed = subprocess.run(
                [
                    executable,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    str(output_dir),
                    str(document.path),
                ],
                check=False,
                capture_output=True,
                text=True,
                timeout=120,
            )
        except subprocess.TimeoutExpired:
            return ExtractedText(
                document_id=document.document_id,
                text="",
                warnings=("estrazione LibreOffice scaduta per timeout",),
            )
        if completed.returncode != 0:
            detail = completed.stderr.strip() or completed.stdout.strip()
            return ExtractedText(
                document_id=document.document_id,
                text="",
                warnings=(f"estrazione LibreOffice non riuscita: {detail}",),
            )
        outputs = sorted(path for path in output_dir.iterdir() if path.is_file())
        if not outputs:
            return ExtractedText(
                document_id=document.document_id,
                text="",
                warnings=("LibreOffice non ha prodotto testo estratto",),
            )
        text = "\n".join(
            path.read_text(encoding="utf-8", errors="replace") for path in outputs
        )
    return ExtractedText(document_id=document.document_id, text=text)


def _extract_pdf_text_with_ocrmypdf(document: DocumentRecord) -> ExtractedText | None:
    """Extract PDF text with system OCRmyPDF sidecar output when available.

    Parameters
    ----------
    document : DocumentRecord
        PDF document.

    Returns
    -------
    ExtractedText | None
        Extracted text, warning from OCRmyPDF, or ``None`` when unavailable.
    """

    executable = shutil.which("ocrmypdf")
    if executable is None:
        return None
    with tempfile.TemporaryDirectory(prefix="sanikey-ocr-") as directory:
        temp_root = Path(directory)
        sidecar = temp_root / "document.txt"
        output_pdf = temp_root / "document.pdf"
        completed = _run_ocrmypdf(
            executable,
            document.path,
            (sidecar, output_pdf),
        )
        if completed.returncode != 0 and _should_retry_ocrmypdf_without_optimize(
            completed
        ):
            sidecar.unlink(missing_ok=True)
            output_pdf.unlink(missing_ok=True)
            retry = _run_ocrmypdf(
                executable,
                document.path,
                (sidecar, output_pdf),
                optimize=False,
            )
            if retry.returncode == 0:
                completed = retry
            else:
                detail = _summarize_command_failure(completed)
                retry_detail = _summarize_command_failure(retry)
                page_detail = _ocrmypdf_failure_page_detail(
                    executable,
                    document.path,
                    optimize=False,
                )
                return ExtractedText(
                    document_id=document.document_id,
                    text="",
                    warnings=(
                        "OCRmyPDF non riuscito; estrazione testo PDF saltata: "
                        f"{detail}; retry con --optimize 0 non riuscito: "
                        f"{retry_detail}{page_detail}",
                    ),
                )
        if completed.returncode != 0:
            detail = _summarize_command_failure(completed)
            page_detail = _ocrmypdf_failure_page_detail(
                executable,
                document.path,
                optimize=True,
            )
            return ExtractedText(
                document_id=document.document_id,
                text="",
                warnings=(
                    "OCRmyPDF non riuscito; estrazione testo PDF saltata: "
                    f"{detail}{page_detail}",
                ),
            )
        if not sidecar.is_file():
            return ExtractedText(
                document_id=document.document_id,
                text="",
                warnings=("OCRmyPDF non ha prodotto il sidecar testuale",),
            )
        return ExtractedText(
            document_id=document.document_id,
            text=sidecar.read_text(encoding="utf-8", errors="replace"),
        )


def _run_ocrmypdf(
    executable: str,
    source: Path,
    outputs: tuple[Path, Path],
    *,
    optimize: bool = True,
    pages: tuple[int, int] | None = None,
) -> subprocess.CompletedProcess[str]:
    """Run OCRmyPDF for one PDF document.

    Parameters
    ----------
    executable : str
        OCRmyPDF executable path.
    source : pathlib.Path
        Source PDF.
    outputs : tuple[pathlib.Path, pathlib.Path]
        Text sidecar and temporary OCR PDF output paths.
    optimize : bool, optional
        Whether OCRmyPDF should run its PDF optimization phase.
    pages : tuple[int, int] | None, optional
        Inclusive one-based page range to process, or ``None`` for all pages.

    Returns
    -------
    subprocess.CompletedProcess[str]
        Completed OCRmyPDF process.
    """

    command = [
        executable,
        "--skip-text",
        "--output-type",
        "pdf",
        "--continue-on-soft-render-error",
    ]
    if not optimize:
        command.extend(("--optimize", "0"))
    if pages is not None:
        start, end = pages
        page_range = str(start) if start == end else f"{start}-{end}"
        command.extend(("--pages", page_range))
    sidecar, output_pdf = outputs
    command.extend(("--sidecar", str(sidecar), str(source), str(output_pdf)))
    return subprocess.run(
        command,
        check=False,
        capture_output=True,
        text=True,
    )


def _ocrmypdf_failure_page_detail(
    executable: str,
    source: Path,
    *,
    optimize: bool,
) -> str:
    """Return a warning suffix that identifies a failing source PDF page.

    Parameters
    ----------
    executable : str
        OCRmyPDF executable path.
    source : pathlib.Path
        Source PDF.
    optimize : bool
        Whether diagnostic OCRmyPDF runs should use optimization.

    Returns
    -------
    str
        Warning suffix containing a one-based source page number, or an empty
        string when the page cannot be determined.
    """

    page_count = _pdf_page_count(source)
    if page_count is None:
        return ""
    page_number = _find_ocrmypdf_failure_page(
        executable,
        source,
        page_count,
        optimize=optimize,
    )
    if page_number is None:
        return ""
    return f"; pagina sorgente non riuscita: {page_number}"


def _pdf_page_count(source: Path) -> int | None:
    """Return the number of pages in a PDF.

    Parameters
    ----------
    source : pathlib.Path
        Source PDF.

    Returns
    -------
    int | None
        Page count, or ``None`` when the PDF cannot be opened.
    """

    try:
        import fitz

        with fitz.open(source) as pdf:
            return len(pdf)
    except (ImportError, RuntimeError, ValueError):
        return None


def _find_ocrmypdf_failure_page(
    executable: str,
    source: Path,
    page_count: int,
    *,
    optimize: bool,
) -> int | None:
    """Find the first OCRmyPDF-failing source page with bisection.

    Parameters
    ----------
    executable : str
        OCRmyPDF executable path.
    source : pathlib.Path
        Source PDF.
    page_count : int
        Number of source PDF pages.
    optimize : bool
        Whether diagnostic OCRmyPDF runs should use optimization.

    Returns
    -------
    int | None
        One-based failing source page, or ``None`` when no failing page can be
        isolated.
    """

    if page_count < 1:
        return None
    start = 1
    end = page_count
    with tempfile.TemporaryDirectory(prefix="sanikey-ocr-diagnose-") as directory:
        temp_root = Path(directory)
        while start < end:
            midpoint = (start + end) // 2
            if _ocrmypdf_page_range_fails(
                executable,
                source,
                temp_root,
                pages=(start, midpoint),
                optimize=optimize,
            ):
                end = midpoint
            else:
                start = midpoint + 1
        if _ocrmypdf_page_range_fails(
            executable,
            source,
            temp_root,
            pages=(start, start),
            optimize=optimize,
        ):
            return start
    return None


def _ocrmypdf_page_range_fails(
    executable: str,
    source: Path,
    temp_root: Path,
    *,
    pages: tuple[int, int],
    optimize: bool,
) -> bool:
    """Return whether OCRmyPDF fails for an inclusive source page range.

    Parameters
    ----------
    executable : str
        OCRmyPDF executable path.
    source : pathlib.Path
        Source PDF.
    temp_root : pathlib.Path
        Temporary directory for diagnostic outputs.
    pages : tuple[int, int]
        Inclusive one-based page range.
    optimize : bool
        Whether OCRmyPDF should run its PDF optimization phase.

    Returns
    -------
    bool
        ``True`` when OCRmyPDF exits with a non-zero status.
    """

    start, end = pages
    sidecar = temp_root / f"pages-{start}-{end}.txt"
    output_pdf = temp_root / f"pages-{start}-{end}.pdf"
    sidecar.unlink(missing_ok=True)
    output_pdf.unlink(missing_ok=True)
    completed = _run_ocrmypdf(
        executable,
        source,
        (sidecar, output_pdf),
        optimize=optimize,
        pages=pages,
    )
    return completed.returncode != 0


def _should_retry_ocrmypdf_without_optimize(
    completed: subprocess.CompletedProcess[str],
) -> bool:
    """Return whether OCRmyPDF should be retried without optimization.

    Parameters
    ----------
    completed : subprocess.CompletedProcess[str]
        Failed OCRmyPDF process.

    Returns
    -------
    bool
        ``True`` when the failure appears to happen after OCR during PDF
        optimization or image transcoding.
    """

    output = f"{completed.stderr}\n{completed.stdout}".lower()
    return any(
        marker in output
        for marker in (
            "optimize_pdf",
            "optimize.py",
            "transcode_jpegs",
            "image file is truncated",
        )
    )


def _summarize_command_failure(completed: subprocess.CompletedProcess[str]) -> str:
    """Summarize a failed external command without dumping full logs.

    Parameters
    ----------
    completed : subprocess.CompletedProcess[str]
        Failed command result.

    Returns
    -------
    str
        Short failure detail suitable for user-facing reports.
    """

    lines = [
        line.strip()
        for text in (completed.stderr, completed.stdout)
        for line in text.splitlines()
        if line.strip()
    ]
    for line in reversed(lines):
        if any(marker in line for marker in ("Error", "Exception", "OSError")):
            return line
    if lines:
        return lines[-1]
    return f"exit status {completed.returncode}"
